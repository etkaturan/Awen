from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x6B)   # headings
TEAL   = RGBColor(0x00, 0x87, 0x8A)   # sub-headings / accents
ORANGE = RGBColor(0xE8, 0x6A, 0x1A)   # call-outs / tips
GRAY   = RGBColor(0x44, 0x44, 0x44)   # body
LIGHT  = RGBColor(0xF2, 0xF6, 0xFA)   # table header fill  (simulated via text)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        val = kwargs.get(edge, {})
        if val:
            el = OxmlElement(f'w:{edge}')
            for k, v in val.items():
                el.set(qn(f'w:{k}'), str(v))
            tcBorders.append(el)
    tcPr.append(tcBorders)

# ── Typography helpers ────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1B3A6B')
    pBdr.append(bot)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.font.name = 'Arial'
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    run.font.name = 'Arial'
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = 'Arial'
    return p

def body(text, bold_parts=None):
    """Add a body paragraph. bold_parts = list of substrings to bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                if idx > 0:
                    r = p.add_run(remaining[:idx])
                    r.font.color.rgb = GRAY; r.font.size = Pt(11); r.font.name='Arial'
                r2 = p.add_run(bp)
                r2.bold = True; r2.font.color.rgb = GRAY; r2.font.size=Pt(11); r2.font.name='Arial'
                remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.color.rgb = GRAY; r.font.size = Pt(11); r.font.name='Arial'
    else:
        run = p.add_run(text)
        run.font.color.rgb = GRAY
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.color.rgb = GRAY
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    return p

def tip(text, label="💡 TIP"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shade_paragraph(p, 'FFF3E0')
    r1 = p.add_run(f"{label}  ")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size=Pt(11); r1.font.name='Arial'
    r2 = p.add_run(text)
    r2.font.color.rgb = GRAY; r2.font.size=Pt(11); r2.font.name='Arial'
    return p

def note(text, label="📌 NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shade_paragraph(p, 'E8F4F8')
    r1 = p.add_run(f"{label}  ")
    r1.bold = True; r1.font.color.rgb = TEAL; r1.font.size=Pt(11); r1.font.name='Arial'
    r2 = p.add_run(text)
    r2.font.color.rgb = GRAY; r2.font.size=Pt(11); r2.font.name='Arial'
    return p

def page_break():
    doc.add_page_break()

# ── Table builder ─────────────────────────────────────────────────────────────
def make_table(headers, rows, col_widths=None):
    """headers = list of str, rows = list of list of str"""
    ncols = len(headers)
    tbl   = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], '1B3A6B')
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.color.rgb = WHITE
        run.font.size = Pt(10); run.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(0)

    # data rows
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        fill  = 'F2F6FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            shade_cell(cells[ci], fill)
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.color.rgb = GRAY
            run.font.size = Pt(10); run.font.name='Arial'
            p.paragraph_format.space_after = Pt(0)

    # column widths
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '1B3A6B')
r = p.add_run("THE COMPLETE PROFESSIONAL")
r.bold=True; r.font.size=Pt(28); r.font.color.rgb=WHITE; r.font.name='Arial'

p2 = doc.add_paragraph()
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p2, '1B3A6B')
r2 = p2.add_run("SOFTWARE DEVELOPER ROADMAP")
r2.bold=True; r2.font.size=Pt(28); r2.font.color.rgb=WHITE; r2.font.name='Arial'

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p3, '00878A')
r3 = p3.add_run("From CS Graduate → Job-Ready Full-Stack Developer → Freelance Professional")
r3.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=WHITE; r3.font.name='Arial'

doc.add_paragraph()
doc.add_paragraph()

for line in ["A Complete Learning Booklet", "Frontend · Backend · DevOps · Cloud · Security · Freelancing",
             "With Free Certifications & Career Strategy"]:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size=Pt(12); r.font.color.rgb=NAVY; r.font.name='Arial'

page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    ("1", "How to Use This Booklet", "3"),
    ("2", "Your Current Skill Map & Gap Analysis", "4"),
    ("3", "Frontend Development — Deep Dive", "5"),
    ("4", "Backend Development — Deep Dive", "10"),
    ("5", "Databases — Relational, NoSQL & Beyond", "15"),
    ("6", "APIs & Communication Protocols", "19"),
    ("7", "Version Control & Collaboration (Git)", "22"),
    ("8", "DevOps, CI/CD & Cloud Hosting", "24"),
    ("9", "Cloud Platforms (AWS, GCP, Azure)", "28"),
    ("10", "Containers & Orchestration (Docker, Kubernetes)", "31"),
    ("11", "Security & Authentication", "33"),
    ("12", "Testing — Unit, Integration & E2E", "36"),
    ("13", "System Design & Architecture", "38"),
    ("14", "Computer Science Fundamentals", "41"),
    ("15", "Soft Skills, Portfolio & Job Strategy", "43"),
    ("16", "Freelancing & Self-Employment Guide", "47"),
    ("17", "Free Certifications Master List", "50"),
    ("18", "Weekly Study Plan (30+ hrs/week)", "53"),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    tab_stops = p.paragraph_format.tab_stops
    from docx.shared import Pt as PT
    r1 = p.add_run(f"{num}.  {title}")
    r1.font.size=Pt(11); r1.font.name='Arial'; r1.font.color.rgb=GRAY
    r2 = p.add_run(f"  ..........  {pg}")
    r2.font.size=Pt(11); r2.font.name='Arial'; r2.font.color.rgb=TEAL

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — HOW TO USE
# ════════════════════════════════════════════════════════════════════════════
h1("1. How to Use This Booklet")
body("This booklet is your complete reference guide for transitioning from a CS graduate with internship experience into a professional, hireable, and eventually self-employed software developer. It is structured in layers:")
bullet("Read Sections 2–14 to understand every domain of modern software development.")
bullet("Use the tables in each section to pick the right tools and frameworks for your goals.")
bullet("Follow Section 18 (Weekly Study Plan) for a 6-month structured schedule.")
bullet("Use Section 17 to collect free certifications that prove your skills to employers.")
bullet("Return to Section 15–16 when applying for jobs or starting freelance work.")
tip("You already know React, TypeScript, Tailwind, Python, SQL, and Vercel — you are NOT starting from zero. This booklet will fill in the professional gaps and expand your toolkit.")
note("Estimated completion time at 30 hrs/week: 5–6 months to be job-ready; 8–10 months to be freelance-ready.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SKILL GAP ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
h1("2. Your Current Skill Map & Gap Analysis")
body("Based on your background, here is an honest assessment of where you stand today and what needs work:")

make_table(
    ["Domain", "What You Know", "What's Missing", "Priority"],
    [
        ["Frontend", "React, TypeScript, Tailwind CSS", "State management (Zustand/Redux), testing, accessibility, performance optimisation", "HIGH"],
        ["Backend", "Python basics, Flask/FastAPI (assumed)", "Node.js/Express, REST API design, auth patterns, background jobs", "HIGH"],
        ["Database", "MySQL, SQLite", "PostgreSQL (industry standard), Redis, MongoDB, ORMs, migrations, indexing", "HIGH"],
        ["Hosting", "Vercel, Umami (analytics)", "Docker, CI/CD pipelines, environment variables, server management, AWS/GCP", "MEDIUM"],
        ["Version Control", "Git basics (assumed)", "Branching strategies, PRs, code review, GitHub Actions", "HIGH"],
        ["Testing", "Unknown", "Unit tests, integration tests, E2E (Playwright/Cypress)", "HIGH"],
        ["Security", "Unknown", "OWASP Top 10, JWT, OAuth2, HTTPS, input validation", "MEDIUM"],
        ["System Design", "Unknown", "Scalability, load balancing, caching, microservices", "MEDIUM"],
        ["DevOps", "Unknown", "Docker, CI/CD, Linux CLI, Nginx", "MEDIUM"],
        ["CS Fundamentals", "University level", "Algorithms/DS practice (LeetCode), Big-O in code reviews", "HIGH"],
    ],
    col_widths=[1.2, 1.5, 2.5, 0.9]
)
tip("Focus on HIGH priority items in months 1-3. MEDIUM items in months 4-6. This is the fastest path to a developer job.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — FRONTEND
# ════════════════════════════════════════════════════════════════════════════
h1("3. Frontend Development — Deep Dive")
body("Frontend development is everything a user sees and interacts with in the browser. It runs on the client side (in the user's browser) and communicates with the backend via HTTP/APIs. You already have a solid foundation here — this section fills in the professional gaps.")

h2("3.1 What Frontend Actually Is")
body("A frontend developer builds user interfaces (UI) using three core browser technologies: HTML (structure), CSS (styling), and JavaScript (interactivity). Modern frontend development adds frameworks, build tools, and performance engineering on top of these foundations.")
bullet("HTML5: semantic elements (<nav>, <main>, <article>), forms, accessibility attributes (ARIA)")
bullet("CSS3: Flexbox, CSS Grid, animations, custom properties (variables), media queries")
bullet("JavaScript (ES2022+): async/await, modules, destructuring, optional chaining, Promises")
note("You already know the framework layer (React + TypeScript + Tailwind). The goal now is to master what's underneath and around it.")

h2("3.2 The React Ecosystem — What You Need to Know Beyond Basics")
body("React is a UI library, not a full framework. Professional React developers are expected to know the entire ecosystem:")

h3("State Management")
body("State management controls how data flows through your application. Choose based on complexity:")
make_table(
    ["Library", "Best For", "Learning Curve", "When to Use"],
    [
        ["useState / useReducer", "Local component state", "You already know this", "Always — start here"],
        ["Context API", "Sharing state between components", "Easy", "Small apps, theming, auth state"],
        ["Zustand", "Global state, simple API", "Very easy", "Medium apps — RECOMMENDED next step"],
        ["Redux Toolkit", "Large apps, time-travel debugging", "Medium", "Large enterprise apps, teams"],
        ["React Query / TanStack Query", "Server state (API data)", "Easy-Medium", "Any app that fetches data — MUST LEARN"],
        ["Jotai", "Atomic state, minimal boilerplate", "Easy", "Alternative to Zustand"],
    ],
    col_widths=[1.5, 1.8, 1.2, 2.0]
)
tip("Learn Zustand + TanStack Query. These two together cover 90% of real-world state needs and are in extremely high demand in job postings.")

h3("Routing")
body("React itself has no routing. These are the main options:")
make_table(
    ["Library", "Used With", "Notes"],
    [
        ["React Router v6", "Any React app", "Industry standard, you likely know it — master nested routes, loaders"],
        ["Next.js App Router", "Next.js only", "File-based routing, server components — learn this FIRST"],
        ["TanStack Router", "Any React app", "Type-safe, growing fast in enterprise"],
    ],
    col_widths=[1.8, 1.5, 3.2]
)

h3("Forms & Validation")
make_table(
    ["Library", "Notes"],
    [
        ["React Hook Form", "Performance-first, minimal re-renders — RECOMMENDED"],
        ["Formik", "Older, verbose — still seen in legacy code"],
        ["Zod", "TypeScript-first schema validation — pair with React Hook Form"],
        ["Yup", "Alternative validation schema library"],
    ],
    col_widths=[2.0, 4.5]
)

h3("UI Component Libraries")
body("You know Tailwind CSS for styling. These libraries give you pre-built accessible components:")
make_table(
    ["Library", "Based On", "Best For", "Notes"],
    [
        ["shadcn/ui", "Radix UI + Tailwind", "Copy-paste components, full control", "MOST POPULAR in 2024-2025"],
        ["Radix UI", "Headless primitives", "Custom-styled accessible components", "shadcn is built on this"],
        ["Chakra UI", "Its own system", "Fast prototyping, good defaults", "Good for freelance projects"],
        ["Headless UI", "Tailwind team", "Simple headless components", "Official Tailwind companion"],
        ["Material UI (MUI)", "Material Design", "Corporate/enterprise look", "Widely used in companies"],
        ["Ant Design", "Ant Group", "Data-heavy admin dashboards", "Very popular in Asian companies"],
    ],
    col_widths=[1.4, 1.3, 1.8, 2.0]
)
tip("Learn shadcn/ui. It's copy-paste, fully customisable with Tailwind (which you know), and is the dominant choice in modern React projects. Employers love seeing it.")

h2("3.3 Next.js — The Framework You Must Learn")
body("Next.js is the most in-demand React framework in the world. It extends React with:")
bullet("Server-Side Rendering (SSR): Pages rendered on the server on each request — great for SEO and dynamic data")
bullet("Static Site Generation (SSG): Pages pre-rendered at build time — extremely fast")
bullet("API Routes: Write backend endpoints inside your frontend project")
bullet("App Router (Next.js 13+): New routing system with React Server Components — learn this, not the old Pages Router")
bullet("Image Optimisation: Automatic image compression and lazy loading")
bullet("Edge Functions: Run code at CDN edge nodes globally")
make_table(
    ["Next.js Concept", "What It Means", "When to Use"],
    [
        ["Server Components", "Render on server, no JS sent to browser", "Default — use for data fetching"],
        ["Client Components", "Classic React with useState, effects", "Add 'use client' — for interactivity"],
        ["Server Actions", "Form submissions handled on server", "Mutations, form handling"],
        ["generateStaticParams", "Pre-render dynamic routes at build", "Blog posts, product pages"],
        ["Middleware", "Run code before request completes", "Auth guards, redirects, A/B testing"],
        ["Route Handlers", "API endpoints (replaces API routes)", "REST endpoints in your Next.js app"],
    ],
    col_widths=[1.8, 2.2, 2.5]
)
note("You are already hosting on Vercel — the company that created Next.js. You are perfectly positioned to learn it. Start here: nextjs.org/learn")

h2("3.4 TypeScript — Going Beyond Basics")
body("You know TypeScript, but professional-level TypeScript goes further:")
bullet("Utility Types: Partial<T>, Required<T>, Pick<T,K>, Omit<T,K>, Record<K,V>, Exclude<T,U>")
bullet("Generic Functions: Writing reusable, type-safe functions with <T>")
bullet("Discriminated Unions: Type narrowing with literal type patterns")
bullet("Type Guards: typeof, instanceof, custom is predicates")
bullet("Zod Integration: Runtime type validation matching your TypeScript types")
bullet("Declaration Files (.d.ts): Writing types for untyped libraries")
tip("One exercise: take a project you built with any, and remove all any types using proper generics and union types. This skill alone impresses interviewers.")

h2("3.5 Performance & Accessibility (Often Ignored, Always Tested in Interviews)")
h3("Performance")
bullet("Core Web Vitals: LCP (Largest Contentful Paint), FID/INP, CLS — Google's ranking metrics")
bullet("Code Splitting: Dynamic imports (next/dynamic, React.lazy) to reduce bundle size")
bullet("Memoisation: useMemo, useCallback, React.memo — when and why (not always needed)")
bullet("Bundle Analysis: Use @next/bundle-analyzer to find large dependencies")
bullet("Image Optimisation: Always use next/image or similar")
h3("Accessibility (a11y)")
bullet("Semantic HTML: Use <button> not <div onClick>, <nav>, <main>, <label for>")
bullet("ARIA attributes: aria-label, aria-describedby, role — when native HTML isn't enough")
bullet("Keyboard navigation: Tab order, focus management, skip links")
bullet("Colour contrast: WCAG AA requires 4.5:1 ratio for normal text")
bullet("Screen reader testing: Use VoiceOver (Mac) or NVDA (Windows) to test")
note("Many government, healthcare, and finance projects legally require WCAG 2.1 AA compliance. This skill gets you into higher-paying contracts.")

h2("3.6 Frontend Testing")
make_table(
    ["Type", "Tool", "What It Tests", "Example"],
    [
        ["Unit", "Vitest / Jest", "Single functions and hooks", "Does formatDate() return the right string?"],
        ["Component", "React Testing Library", "Component output and behaviour", "Does the button show 'Loading' when clicked?"],
        ["E2E", "Playwright (recommended)", "Full user journeys in real browser", "Can a user sign up, log in, and check out?"],
        ["Visual", "Storybook + Chromatic", "UI component appearance", "Did the button change visually?"],
    ],
    col_widths=[1.0, 1.8, 2.0, 2.7]
)
tip("Start with React Testing Library for components and Playwright for E2E. These are the most in-demand in job postings today.")

h2("3.7 Other Frontend Frameworks to Know About")
make_table(
    ["Framework", "Language", "Strengths", "Market Position"],
    [
        ["Vue.js 3", "JavaScript/TS", "Gentle learning curve, great docs, Options & Composition API", "Very popular in Europe and Asia"],
        ["Svelte / SvelteKit", "JavaScript/TS", "No virtual DOM, tiny bundles, elegant syntax", "Growing fast, loved by devs"],
        ["Angular", "TypeScript", "Full opinionated framework, strong DI system", "Dominant in enterprises, Java shops"],
        ["Astro", "JavaScript/TS", "Content sites, ships zero JS by default, island architecture", "Best for blogs, marketing sites"],
        ["Remix", "TypeScript", "Server-first, nested routes, great forms", "Meta now maintains it with React Router"],
        ["Solid.js", "JavaScript/TS", "React-like but faster, fine-grained reactivity", "Niche but growing"],
    ],
    col_widths=[1.4, 1.2, 2.5, 2.0]
)
body("You do not need to learn all of these. React/Next.js covers 60%+ of the job market. Vue is worth a glance if you want European/Asian clients.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — BACKEND
# ════════════════════════════════════════════════════════════════════════════
h1("4. Backend Development — Deep Dive")
body("The backend is the server-side of an application. It handles business logic, authentication, data storage, external integrations, and serves data to the frontend via APIs. You know Python — that gives you a great foundation.")

h2("4.1 What a Backend Developer Does")
bullet("Design and build REST or GraphQL APIs")
bullet("Manage databases: queries, migrations, indexes, relationships")
bullet("Handle authentication and authorisation (who can log in, who can do what)")
bullet("Write business logic: pricing calculations, order processing, notifications")
bullet("Integrate third-party services: payment processors, email, SMS, cloud storage")
bullet("Ensure performance: caching, query optimisation, async processing")
bullet("Secure the application: input validation, SQL injection prevention, rate limiting")

h2("4.2 Python Backend Frameworks")
body("You know Python. Here are the frameworks you should use professionally:")
make_table(
    ["Framework", "Style", "Best For", "Key Features", "Job Demand"],
    [
        ["FastAPI", "Async, modern", "APIs, microservices, ML backends", "Auto docs (Swagger), Pydantic validation, async, type hints", "Very High ↑"],
        ["Django", "Full-stack, batteries included", "Large web apps, admin panels, e-commerce", "ORM, admin interface, auth system, REST framework", "Very High ↑"],
        ["Django REST Framework", "Add-on for Django", "REST APIs in Django apps", "Serialisers, viewsets, permissions, throttling", "Very High ↑"],
        ["Flask", "Minimal/micro", "Small APIs, prototypes, learning", "Lightweight, flexible, you choose everything", "High"],
        ["Litestar", "Async, modern", "High-performance APIs", "FastAPI alternative, OpenAPI, validation", "Growing"],
    ],
    col_widths=[1.3, 1.2, 1.8, 2.2, 0.9]
)
tip("Learn FastAPI + SQLAlchemy + Alembic as your primary Python stack. Then learn Django for when employers need an admin panel or full web app. This combination covers 80% of Python backend jobs.")

h2("4.3 Node.js / JavaScript Backend")
body("Even as a Python developer, knowing Node.js is a massive advantage. Many companies use it exclusively, and many full-stack roles require it.")
make_table(
    ["Framework", "Style", "Best For", "Notes"],
    [
        ["Express.js", "Minimal, unopinionated", "APIs, learning backend fundamentals", "The jQuery of Node — everywhere, foundational"],
        ["Fastify", "Fast, schema-driven", "High-performance APIs", "Drop-in Express alternative, better performance"],
        ["NestJS", "Angular-inspired, TypeScript", "Enterprise APIs, large codebases", "Decorators, DI, modules — very structured"],
        ["Hono", "Ultra-light, edge-ready", "Edge functions, Cloudflare Workers", "Modern, tiny, excellent DX"],
        ["tRPC", "Type-safe RPC", "Full-stack TypeScript with Next.js", "No REST/GraphQL needed, end-to-end types"],
    ],
    col_widths=[1.3, 1.4, 1.8, 2.9]
)
note("Since you already know TypeScript and React, learning Express.js or Fastify + NestJS is a natural extension. tRPC is a game-changer for full-stack TypeScript projects.")

h2("4.4 Other Backend Languages Worth Knowing")
make_table(
    ["Language", "Primary Framework", "Known For", "Learn If..."],
    [
        ["Go (Golang)", "Gin, Echo, Fiber", "Extreme performance, concurrency, microservices", "You want DevOps/backend infrastructure roles"],
        ["Java", "Spring Boot", "Enterprise, banking, Android", "You want enterprise or fintech jobs"],
        ["C# (.NET)", "ASP.NET Core", "Microsoft ecosystem, game backends", "You want to work at Microsoft-stack companies"],
        ["Ruby", "Ruby on Rails", "Rapid prototyping, startup culture", "Less common now but still used at some startups"],
        ["PHP", "Laravel", "WordPress, legacy web apps, still widely used", "You want to do WordPress freelance work"],
        ["Rust", "Actix, Axum", "Systems programming, WebAssembly", "Advanced — performance-critical systems"],
    ],
    col_widths=[1.2, 1.5, 2.3, 2.5]
)
body("You do not need to learn all of these. Python + Node.js covers most job opportunities. Go is worth learning in year 2 if you want senior backend roles.")

h2("4.5 Core Backend Concepts Every Professional Must Know")
h3("REST API Design Principles")
bullet("Resources as nouns: /users, /products, /orders — not /getUser or /createOrder")
bullet("HTTP verbs: GET (read), POST (create), PUT/PATCH (update), DELETE (remove)")
bullet("Status codes: 200 OK, 201 Created, 400 Bad Request, 401 Unauthorised, 403 Forbidden, 404 Not Found, 500 Server Error")
bullet("Versioning: /api/v1/users — always version your APIs")
bullet("Pagination: limit/offset or cursor-based for large datasets")
bullet("HATEOAS: Hypermedia links in responses (advanced but good to know)")

h3("Middleware & Request Lifecycle")
body("In web frameworks, middleware are functions that run before your route handler. Learn to write middleware for:")
bullet("Authentication: Verify JWT token before allowing access")
bullet("Logging: Log every request with method, path, status, duration")
bullet("Rate limiting: Block clients making too many requests")
bullet("CORS: Control which domains can call your API")
bullet("Request validation: Validate body/params before they hit business logic")

h3("Async Programming (Critical)")
body("Modern backend applications are inherently concurrent. Python and Node.js both support async programming:")
bullet("Python: async/await with asyncio, async database drivers (asyncpg, SQLAlchemy async)")
bullet("Python: Celery for background job queues (send emails, process images, scheduled tasks)")
bullet("Node.js: Promises, async/await, EventEmitter, streams")
bullet("Queues: Redis + Bull (Node), Celery + Redis (Python) for background processing")
bullet("WebSockets: Real-time bidirectional communication (Socket.io, ws library)")

h3("Caching Strategies")
make_table(
    ["Strategy", "Tool", "What It Caches", "Example"],
    [
        ["In-memory (app level)", "Python dict, Node Map", "Computed values for current process", "Store parsed config"],
        ["Distributed cache", "Redis", "Shared data across multiple servers", "User sessions, API responses"],
        ["HTTP cache", "Cache-Control headers", "Browser and CDN caching", "Static assets, GET responses"],
        ["Database query cache", "PostgreSQL, application", "Expensive DB query results", "Dashboard stats"],
        ["CDN", "Cloudflare, AWS CloudFront", "Static files globally", "Images, JS bundles, CSS"],
    ],
    col_widths=[1.6, 1.5, 2.0, 2.4]
)
tip("Redis is the single most important tool to learn after your primary language/framework. It is used everywhere: sessions, caching, queues, pub/sub.")

h2("4.6 Serverless & Edge Computing")
body("Serverless does not mean no servers — it means you write functions, and the platform manages the servers, scaling, and billing.")
make_table(
    ["Platform", "Runtime", "Best For", "Free Tier"],
    [
        ["Vercel Functions", "Node.js, Edge", "Next.js API routes, serverless APIs", "Generous free tier"],
        ["Netlify Functions", "Node.js, Deno", "JAMstack backends", "Free tier available"],
        ["AWS Lambda", "Any language", "Event-driven functions, scheduled jobs", "1M requests/month free"],
        ["Cloudflare Workers", "JavaScript (V8)", "Edge computing, ultra-low latency", "100k requests/day free"],
        ["Supabase Edge Functions", "Deno", "Supabase-integrated serverless", "Free tier"],
        ["Google Cloud Functions", "Node, Python, Go, Java", "GCP ecosystem", "2M requests/month free"],
    ],
    col_widths=[1.7, 1.4, 2.0, 1.8]
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATABASES
# ════════════════════════════════════════════════════════════════════════════
h1("5. Databases — Relational, NoSQL & Beyond")
body("You know MySQL and SQLite. Databases are the backbone of every application. Professional developers must understand not just SQL, but when to use which type of database and how to design schemas that scale.")

h2("5.1 Types of Databases")
make_table(
    ["Type", "Examples", "Data Model", "Best For"],
    [
        ["Relational (SQL)", "PostgreSQL, MySQL, SQLite, SQL Server", "Tables, rows, foreign keys, joins", "Structured data, complex relationships, transactions"],
        ["Document", "MongoDB, CouchDB, Firestore", "JSON-like documents in collections", "Flexible schemas, content management, catalogs"],
        ["Key-Value", "Redis, DynamoDB", "Key → Value pairs", "Caching, sessions, real-time leaderboards"],
        ["Column-Family", "Apache Cassandra, HBase", "Rows with dynamic column families", "Write-heavy, time-series, IoT data at scale"],
        ["Graph", "Neo4j, Amazon Neptune", "Nodes and edges (relationships)", "Social networks, recommendation engines, fraud detection"],
        ["Search", "Elasticsearch, Meilisearch, Typesense", "Inverted index over documents", "Full-text search, log analysis, autocomplete"],
        ["Time-Series", "InfluxDB, TimescaleDB", "Timestamps + measurements", "Monitoring, analytics, sensor data"],
        ["Vector", "Pinecone, Qdrant, pgvector", "High-dimensional embeddings", "AI/ML semantic search, recommendation systems"],
    ],
    col_widths=[1.4, 2.0, 1.8, 2.3]
)

h2("5.2 PostgreSQL — Your New Primary Database")
body("PostgreSQL (Postgres) is the most feature-rich, reliable, open-source relational database. It should replace MySQL as your default choice because:")
bullet("Advanced data types: JSON/JSONB, arrays, UUIDs, ranges, enums, composite types")
bullet("Full-text search: Built-in without Elasticsearch for many use cases")
bullet("Window functions: Powerful analytics queries (RANK, ROW_NUMBER, LAG, LEAD)")
bullet("CTEs (Common Table Expressions): Write readable, recursive queries with WITH clauses")
bullet("Extensions: pgvector (AI embeddings), PostGIS (geolocation), pg_trgm (fuzzy search)")
bullet("ACID compliance: Rock-solid transactions with real-world reliability")
bullet("Industry adoption: Supabase, Heroku, Railway, Render, AWS RDS all default to Postgres")
note("Every major cloud platform, startup toolkit, and modern SaaS is built on PostgreSQL. MySQL is fine but learning Postgres sets you apart.")

h2("5.3 SQL Mastery — What Professionals Know")
h3("Basic to Intermediate (You should know these)")
bullet("SELECT, WHERE, ORDER BY, LIMIT, OFFSET, DISTINCT")
bullet("JOINs: INNER, LEFT, RIGHT, FULL OUTER, CROSS")
bullet("Aggregates: COUNT, SUM, AVG, MIN, MAX with GROUP BY and HAVING")
bullet("Subqueries: Correlated and non-correlated")
bullet("String functions: LIKE, ILIKE, CONCAT, SUBSTRING, TRIM")
bullet("Date functions: NOW(), DATE_TRUNC, EXTRACT, AGE, INTERVAL")

h3("Advanced SQL (Learn This to Stand Out)")
bullet("Window functions: SELECT name, salary, RANK() OVER (PARTITION BY dept ORDER BY salary DESC)")
bullet("CTEs: WITH cte AS (SELECT ...) SELECT * FROM cte WHERE ...")
bullet("Recursive CTEs: For hierarchical data (org charts, categories, file systems)")
bullet("EXPLAIN ANALYZE: Understand query execution plans, identify slow queries")
bullet("Indexes: B-tree, Hash, GiST, GIN — when to create them, composite indexes")
bullet("Transactions and isolation levels: READ COMMITTED, REPEATABLE READ, SERIALIZABLE")
bullet("Stored procedures and triggers: When to use them (rarely, but you must know)")
bullet("JSON queries: ->, ->>, #>, jsonb_path_query for querying JSON columns")

h2("5.4 ORMs — Object-Relational Mappers")
body("ORMs let you interact with databases using your programming language instead of raw SQL. Professional developers know when to use them and when to bypass them:")
make_table(
    ["ORM", "Language", "Style", "Notes"],
    [
        ["SQLAlchemy 2.0", "Python", "Data Mapper pattern, async support", "Most powerful Python ORM — MUST LEARN for Python backend"],
        ["Alembic", "Python", "Migration tool for SQLAlchemy", "Manage schema changes — pair with SQLAlchemy"],
        ["Django ORM", "Python", "Active Record, built into Django", "Easy, powerful for Django apps"],
        ["Prisma", "TypeScript/JS", "Type-safe, code-first schema", "MUST LEARN for Node.js — generates perfect TypeScript types"],
        ["Drizzle ORM", "TypeScript/JS", "SQL-like, lightweight, type-safe", "Growing fast, great for serverless"],
        ["Sequelize", "JavaScript", "Active Record, callbacks-era", "Older, still widely used in legacy Node apps"],
        ["TypeORM", "TypeScript", "Decorators, Active Record/Data Mapper", "Popular in NestJS apps"],
    ],
    col_widths=[1.4, 1.2, 2.0, 2.9]
)
tip("Python path: SQLAlchemy + Alembic. Node.js path: Prisma. These are the market leaders. For freelance work with Next.js: Prisma + PostgreSQL (on Supabase or Neon) is the dominant stack.")

h2("5.5 Redis — The Swiss Army Knife")
body("Redis is an in-memory data structure store. It is NOT just a cache — it is a full-featured data platform:")
bullet("String caching: Store API responses, computed values with TTL (time-to-live)")
bullet("Sessions: Store user session data server-side")
bullet("Queues: Bull/BullMQ (Node.js), Celery (Python) use Redis as a message broker")
bullet("Pub/Sub: Broadcast messages to multiple subscribers (real-time notifications)")
bullet("Rate limiting: Track request counts per user/IP with INCR and EXPIRE")
bullet("Leaderboards: Sorted sets (ZADD, ZRANK) for real-time ranking systems")
bullet("Distributed locks: Prevent race conditions in distributed systems")
note("Redis is used in virtually every production application at scale. It is also a common interview topic. Learn it with the ioredis (Node) or redis-py (Python) library.")

h2("5.6 MongoDB — Document Database")
body("MongoDB stores JSON-like documents (BSON). It is popular for:")
bullet("Content management systems with flexible, varying schemas")
bullet("Rapid prototyping where schema changes frequently")
bullet("Storing hierarchical or nested data naturally")
bullet("Node.js + Express + MongoDB + React (MERN stack) — huge in the freelance market")
make_table(
    ["MongoDB Concept", "SQL Equivalent", "Notes"],
    [
        ["Collection", "Table", "Group of related documents"],
        ["Document", "Row", "JSON object with any structure"],
        ["Field", "Column", "Can contain nested objects or arrays"],
        ["_id", "Primary Key", "Auto-generated ObjectId or custom"],
        ["Aggregation Pipeline", "GROUP BY + JOINs", "Powerful data transformation stages ($match, $group, $lookup)"],
        ["Atlas", "Cloud RDS", "MongoDB's managed cloud service — free 512MB tier"],
    ],
    col_widths=[1.8, 1.5, 3.2]
)

h2("5.7 Supabase — The Postgres Platform to Learn Now")
body("Supabase is an open-source Firebase alternative built on PostgreSQL. It is extremely popular for solo developers and freelancers:")
bullet("Hosted PostgreSQL database with a visual table editor")
bullet("Auto-generated REST and GraphQL APIs from your database schema")
bullet("Built-in authentication (email, OAuth, magic links, phone)")
bullet("Real-time subscriptions via WebSockets")
bullet("Storage: File uploads with S3-compatible bucket API")
bullet("Edge Functions: Deno serverless functions")
tip("Supabase + Next.js is one of the most popular freelance stacks in 2024-2025. You can build a full SaaS product with zero server management. Since you already use Vercel, adding Supabase is a natural extension.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — APIs
# ════════════════════════════════════════════════════════════════════════════
h1("6. APIs & Communication Protocols")
body("APIs (Application Programming Interfaces) are how software systems communicate. As a full-stack developer, you will both consume APIs (in frontend) and build them (in backend).")

h2("6.1 REST — The Standard")
body("REST (Representational State Transfer) is the dominant API architecture. You likely use it already. Professional-level REST knowledge includes:")
make_table(
    ["Concept", "Details"],
    [
        ["Resource Naming", "Plural nouns: /users, /products/123/reviews — never verbs like /getUser"],
        ["HTTP Methods", "GET (read), POST (create), PUT (replace), PATCH (partial update), DELETE"],
        ["Status Codes", "200, 201, 204, 400, 401, 403, 404, 409, 422, 429, 500, 503"],
        ["Headers", "Content-Type, Accept, Authorization, ETag, Cache-Control, X-Request-ID"],
        ["Pagination", "?page=1&limit=20 (offset) or ?cursor=xyz (cursor-based for large datasets)"],
        ["Filtering & Sorting", "?status=active&sort=created_at&order=desc"],
        ["Versioning", "/api/v1/ prefix or Accept-Version header"],
        ["Error Responses", "Consistent JSON: { error: 'message', code: 'ERROR_CODE', details: {...} }"],
        ["HATEOAS", "Include links to related resources in responses (advanced)"],
        ["OpenAPI/Swagger", "Document your API with a spec file — auto-generates docs"],
    ],
    col_widths=[2.0, 4.5]
)

h2("6.2 GraphQL")
body("GraphQL is a query language for APIs that lets clients request exactly the data they need — no more, no less. Created by Meta (Facebook), it is extremely popular in startups and product companies.")
bullet("Schema-first: Define your API types in a schema definition language (SDL)")
bullet("Single endpoint: All requests go to /graphql — clients specify what they need")
bullet("No over/under-fetching: Get exactly the fields you ask for")
bullet("Subscriptions: Real-time data via WebSockets using the GraphQL protocol")
bullet("Introspection: Clients can query the schema itself to discover available types")
make_table(
    ["Tool", "Type", "Notes"],
    [
        ["Apollo Server", "Node.js GraphQL server", "Most popular, rich ecosystem"],
        ["Strawberry", "Python GraphQL", "Pythonic, type-hint based — RECOMMENDED for Python"],
        ["Ariadne", "Python GraphQL", "Schema-first Python library"],
        ["Apollo Client", "Frontend GraphQL client", "Caching, optimistic updates, React integration"],
        ["urql", "Frontend GraphQL client", "Lighter alternative to Apollo Client"],
        ["GraphQL Yoga", "Node.js server", "Modern, spec-compliant, works everywhere"],
    ],
    col_widths=[1.8, 1.5, 3.2]
)
note("GraphQL is not better than REST — it solves different problems. REST is simpler and more cacheable. GraphQL shines in complex data requirements and mobile apps where bandwidth matters. Learn REST first, GraphQL second.")

h2("6.3 WebSockets & Real-Time Communication")
make_table(
    ["Technology", "Description", "Use Cases"],
    [
        ["WebSockets", "Persistent bidirectional connection between client and server", "Chat, live notifications, collaborative editing, dashboards"],
        ["Socket.io", "WebSocket library with fallback and rooms support", "Chat apps, multiplayer games, real-time dashboards"],
        ["Server-Sent Events (SSE)", "One-way server → client streaming", "Live feeds, progress bars, AI streaming (ChatGPT-style)"],
        ["Long Polling", "Client repeatedly asks server for updates", "Simple real-time, fallback when WebSockets not available"],
        ["WebRTC", "Peer-to-peer media/data between browsers", "Video calls, screen sharing, file transfer"],
    ],
    col_widths=[1.6, 2.2, 2.7]
)
tip("SSE (Server-Sent Events) is underrated and easy to implement. It is how most AI streaming responses work. Learn it alongside WebSockets.")

h2("6.4 gRPC — The Backend-to-Backend Protocol")
body("gRPC (Google Remote Procedure Call) is a high-performance RPC framework that uses Protocol Buffers (protobuf) for serialisation. It is used between microservices where performance matters. You don't need this for job searching now, but knowing it exists and roughly how it works is useful for system design interviews.")
bullet("Binary protocol: Much faster than JSON for service-to-service calls")
bullet("Strongly typed: Protobuf schema enforces contract between services")
bullet("Streaming: Client streaming, server streaming, and bidirectional streaming")
bullet("Code generation: Generate client/server code in any language from the .proto file")

h2("6.5 Third-Party API Integrations (Freelance Gold)")
body("Most freelance projects involve integrating third-party services. Here are the most common and most lucrative:")
make_table(
    ["Category", "Services", "Why It Matters"],
    [
        ["Payments", "Stripe, PayPal, Paddle, Lemon Squeezy", "Every e-commerce project needs this — Stripe is the gold standard"],
        ["Email", "Resend, SendGrid, Postmark, Mailgun", "Transactional emails (receipts, OTPs, welcome emails)"],
        ["SMS", "Twilio, Vonage", "OTP verification, alerts"],
        ["Auth", "Auth0, Clerk, Firebase Auth, NextAuth.js", "Social login (Google, GitHub), enterprise SSO"],
        ["Storage", "AWS S3, Cloudflare R2, Supabase Storage", "File uploads, images, documents"],
        ["Maps", "Google Maps, Mapbox, Leaflet", "Location features"],
        ["AI / LLMs", "OpenAI, Anthropic, Google Gemini, Replicate", "AI features — growing extremely fast"],
        ["Search", "Algolia, Meilisearch, Typesense", "Instant search in apps"],
        ["Notifications", "Novu, Knock, OneSignal", "Push, email, in-app notifications"],
    ],
    col_widths=[1.3, 2.5, 2.7]
)
tip("Learn Stripe properly. Stripe integration is one of the most requested and well-paid freelance tasks. Their documentation is excellent.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — GIT
# ════════════════════════════════════════════════════════════════════════════
h1("7. Version Control & Collaboration (Git)")
body("Git is non-negotiable. Every professional developer uses it daily. Beyond basic commits and pushes, you need to know collaborative Git workflows.")

h2("7.1 Git Commands Every Professional Uses")
make_table(
    ["Command", "What It Does"],
    [
        ["git commit --amend", "Fix the last commit message or add forgotten files"],
        ["git rebase -i HEAD~3", "Interactively squash, reorder, or edit the last 3 commits"],
        ["git stash / git stash pop", "Temporarily save uncommitted changes"],
        ["git cherry-pick <hash>", "Apply a specific commit from another branch"],
        ["git bisect", "Binary search through commits to find when a bug was introduced"],
        ["git log --oneline --graph", "Visualise branch history in the terminal"],
        ["git reset --soft HEAD~1", "Undo last commit but keep changes staged"],
        ["git reset --hard HEAD~1", "Undo last commit and discard all changes (CAREFUL)"],
        ["git blame <file>", "See who wrote each line of a file"],
        ["git reflog", "See ALL recent HEAD positions — recover from mistakes"],
    ],
    col_widths=[2.5, 4.0]
)

h2("7.2 Branching Strategies")
make_table(
    ["Strategy", "Branch Structure", "Best For"],
    [
        ["GitHub Flow", "main + feature branches + PRs", "Continuous deployment, small teams, most modern startups"],
        ["Git Flow", "main + develop + feature + release + hotfix", "Versioned releases, large teams"],
        ["Trunk-Based Development", "main + short-lived feature branches", "High-frequency deployments, large tech companies"],
    ],
    col_widths=[1.8, 2.5, 2.2]
)
tip("Learn GitHub Flow — it is what 90% of startups and modern companies use. Create a branch, make commits, open a PR, get reviewed, merge to main, deploy.")

h2("7.3 Pull Requests & Code Review")
body("In professional teams, no code goes to main without a Pull Request (PR) and code review. How to write good PRs:")
bullet("Small PRs: Under 400 lines of diff — reviewers won't thoroughly review massive PRs")
bullet("Good title: '[Feature] Add user profile editing' — not just 'updates'")
bullet("Description: What does this PR do? Why? How to test it? Screenshots for UI changes?")
bullet("Self-review first: Review your own PR before requesting others")
bullet("Link to issue: 'Closes #123' automatically closes the linked GitHub issue on merge")

h2("7.4 GitHub Features for Your Portfolio")
bullet("GitHub Actions: CI/CD pipelines that run on every push (tests, linting, deployment)")
bullet("GitHub Pages: Host static sites for free from a repository")
bullet("GitHub Projects: Kanban boards for managing your learning or freelance projects")
bullet("README badges: Show build status, test coverage, npm version in your repo")
bullet("Pinned repositories: Curate your best 6 projects on your GitHub profile")
note("Your GitHub profile is your public CV. Employers look at it. Green contribution squares every day (or most days) signal that you code consistently.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — DEVOPS
# ════════════════════════════════════════════════════════════════════════════
h1("8. DevOps, CI/CD & Cloud Hosting")
body("DevOps bridges the gap between development and operations. You don't need to be a DevOps engineer to get a developer job, but you must understand the concepts to work effectively on any professional team.")

h2("8.1 Linux CLI — The Foundation")
body("Every server runs Linux. You must be comfortable with the command line:")
make_table(
    ["Command Category", "Examples", "Why It Matters"],
    [
        ["Navigation", "ls, cd, pwd, find, locate", "Move around file systems"],
        ["File operations", "cp, mv, rm, mkdir, touch, cat, less, tail -f", "Manage files and logs"],
        ["Permissions", "chmod, chown, ls -la, umask", "Security and file access control"],
        ["Process management", "ps aux, kill, htop, systemctl, journalctl", "Monitor and control running services"],
        ["Networking", "curl, wget, netstat, ss, ping, nslookup, dig", "Debug HTTP and network issues"],
        ["Text processing", "grep, awk, sed, sort, uniq, wc, cut", "Parse logs and process data"],
        ["SSH", "ssh, scp, rsync, ssh-keygen", "Connect to remote servers securely"],
        ["Package managers", "apt, yum, brew, snap", "Install software on servers"],
        ["Environment", "export, env, .bashrc, .env files, source", "Manage environment variables"],
        ["Cron jobs", "crontab -e, cron syntax (*/5 * * * *)", "Schedule recurring tasks"],
    ],
    col_widths=[1.6, 2.2, 2.7]
)
tip("Install WSL2 (Windows Subsystem for Linux) or use macOS Terminal. Practice Linux every day — even just SSHing into a cheap VPS ($5/month DigitalOcean) and deploying a project teaches you more than any tutorial.")

h2("8.2 CI/CD Pipelines")
body("CI/CD stands for Continuous Integration / Continuous Delivery (or Deployment). It automates testing, building, and deploying your code on every push.")
make_table(
    ["Stage", "What Happens", "Tools"],
    [
        ["Trigger", "Developer pushes to a branch or opens a PR", "Git + GitHub/GitLab"],
        ["CI: Lint", "Check code style (ESLint, Black, Ruff)", "GitHub Actions, GitLab CI"],
        ["CI: Test", "Run all automated tests", "Vitest, pytest, Jest"],
        ["CI: Build", "Compile TypeScript, build Docker image", "esbuild, webpack, Docker"],
        ["CI: Security scan", "Check for vulnerable dependencies", "Dependabot, Snyk, Trivy"],
        ["CD: Deploy to staging", "Deploy to a test environment", "Vercel Preview, Render"],
        ["CD: Deploy to production", "After merge to main, deploy live", "Vercel, AWS, Railway"],
        ["Post-deploy: Monitor", "Check error rates and performance", "Sentry, Datadog"],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

h3("GitHub Actions (Start Here)")
body("GitHub Actions is the easiest CI/CD tool to learn because it is built into GitHub, where you already keep your code. A basic workflow file:")
bullet("Stored in .github/workflows/ci.yml in your repository")
bullet("Triggered on push to main, pull requests, or schedules")
bullet("Runs on GitHub's cloud runners (Linux, Windows, macOS)")
bullet("Free for public repos; 2000 minutes/month free for private repos")

h3("Other CI/CD Tools")
make_table(
    ["Tool", "Best For", "Notes"],
    [
        ["GitHub Actions", "Any project on GitHub", "Start here — easiest, tightly integrated"],
        ["GitLab CI/CD", "GitLab users", "Very powerful, YAML-based pipelines"],
        ["CircleCI", "Complex pipelines, parallel jobs", "Popular in established companies"],
        ["Jenkins", "Self-hosted, legacy systems", "Very customisable, but complex to maintain"],
        ["Vercel / Netlify", "Frontend auto-deploy", "Zero-config CD — push to main = deployed"],
    ],
    col_widths=[1.5, 1.8, 3.2]
)

h2("8.3 Hosting Platforms — Beyond Vercel")
body("You know Vercel. Here is the full landscape of hosting options:")
make_table(
    ["Platform", "Type", "Best For", "Pricing Model", "Notes"],
    [
        ["Vercel", "PaaS / Serverless", "Next.js, React frontends", "Free + usage", "You already know this — excellent"],
        ["Netlify", "PaaS / Serverless", "JAMstack, static sites", "Free + usage", "Vercel competitor"],
        ["Railway", "PaaS", "Full-stack apps, backends, databases", "Usage-based", "Easiest way to deploy backends"],
        ["Render", "PaaS", "Web services, APIs, databases", "Free + paid tiers", "Heroku alternative, generous free tier"],
        ["Fly.io", "Container PaaS", "Dockerised apps, global edge", "Usage-based", "Run Docker containers cheaply"],
        ["DigitalOcean", "Cloud VPS + PaaS", "VPS control, managed databases", "From $5/month", "Great for learning server management"],
        ["Hetzner", "Cloud VPS", "Cheap European servers", "From €3.29/month", "Best price/performance for VPS"],
        ["AWS", "Full cloud", "Enterprise, scalable infrastructure", "Pay-per-use", "Industry dominant — learn the basics"],
        ["Google Cloud", "Full cloud", "ML/AI, Firebase, GKE", "Pay-per-use", "Best for AI workloads"],
        ["Supabase", "BaaS", "Postgres + auth + storage + functions", "Free + paid", "Full backend in one"],
    ],
    col_widths=[1.2, 1.3, 1.8, 1.2, 1.9]
)
tip("For freelance work: Vercel (frontend) + Railway or Render (backend) + Supabase (database) is a zero-config, low-cost, high-productivity stack. You can deploy a full SaaS in an afternoon.")

h2("8.4 Nginx — The Web Server You Must Know")
body("Nginx (pronounced 'engine-x') is the most widely used web server. In production, it often sits in front of your application server:")
bullet("Reverse proxy: Forward requests from the internet to your app (running on a local port)")
bullet("SSL/TLS termination: Handle HTTPS certificates so your app only deals with HTTP internally")
bullet("Static file serving: Serve static assets directly without hitting your app server")
bullet("Load balancing: Distribute traffic across multiple app server instances")
bullet("Rate limiting: Block abusive clients at the server level")
bullet("Gzip compression: Compress responses before sending to reduce bandwidth")
note("You don't need to master Nginx for a developer job, but knowing how to write a basic server block and set up SSL (via Certbot/Let's Encrypt) makes you capable of managing your own VPS.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — CLOUD
# ════════════════════════════════════════════════════════════════════════════
h1("9. Cloud Platforms (AWS, GCP, Azure)")
body("Cloud platforms provide on-demand computing resources over the internet. Rather than buying and managing physical servers, you rent capacity from a cloud provider and pay for what you use. AWS dominates the market with ~33% share, followed by Azure (~22%) and Google Cloud (~11%).")

h2("9.1 AWS — The Market Leader")
body("Even if you never become an AWS engineer, knowing the core services is essential for job interviews and working in any tech company.")
make_table(
    ["Service", "Category", "What It Does", "Learn Priority"],
    [
        ["EC2", "Compute", "Virtual servers (VMs) you fully control", "MEDIUM"],
        ["Lambda", "Serverless", "Run functions without managing servers", "HIGH"],
        ["S3", "Storage", "Object storage for files, images, backups, static sites", "HIGH"],
        ["RDS", "Database", "Managed relational databases (Postgres, MySQL, etc.)", "HIGH"],
        ["DynamoDB", "Database", "Serverless NoSQL key-value and document database", "MEDIUM"],
        ["ElastiCache", "Cache", "Managed Redis and Memcached", "MEDIUM"],
        ["CloudFront", "CDN", "Global content delivery network, caches S3 content", "HIGH"],
        ["Route 53", "DNS", "Domain registration and DNS management", "MEDIUM"],
        ["IAM", "Security", "Identity and Access Management — users, roles, permissions", "HIGH"],
        ["VPC", "Networking", "Virtual Private Cloud — your isolated network in AWS", "MEDIUM"],
        ["ECS / EKS", "Containers", "Run Docker containers (ECS=simpler, EKS=Kubernetes)", "MEDIUM"],
        ["Cognito", "Auth", "User authentication and authorisation service", "LOW"],
        ["SES", "Email", "Simple Email Service for transactional emails", "MEDIUM"],
        ["SQS/SNS", "Messaging", "Queue (SQS) and notification (SNS) services", "MEDIUM"],
        ["CloudWatch", "Monitoring", "Logs, metrics, alarms for your AWS resources", "HIGH"],
    ],
    col_widths=[1.3, 1.2, 2.8, 1.1]
)
tip("The AWS Free Tier gives you 12 months of free access to many services. Create a free account and deploy a simple project: Lambda + API Gateway + S3 + DynamoDB. This teaches you the fundamentals and looks great on a CV.")

h2("9.2 The AWS Certification Path")
make_table(
    ["Certification", "Level", "Cost", "Focus"],
    [
        ["AWS Cloud Practitioner", "Foundational", "$100", "Cloud concepts, billing, services overview — great first cert"],
        ["AWS Solutions Architect Associate", "Associate", "$150", "Designing cloud architectures — most popular AWS cert"],
        ["AWS Developer Associate", "Associate", "$150", "Building and deploying applications on AWS"],
        ["AWS SysOps Administrator", "Associate", "$150", "Operations, monitoring, deployments"],
        ["AWS DevOps Professional", "Professional", "$300", "Advanced CI/CD, microservices, automation"],
    ],
    col_widths=[2.5, 1.2, 0.8, 3.0]
)
note("The AWS Cloud Practitioner exam can be prepared for in 2-4 weeks with free resources (freeCodeCamp YouTube, AWS Skill Builder). It is a strong signal on a CV for junior-mid developer roles.")

h2("9.3 Google Cloud Platform (GCP)")
body("GCP is particularly strong for AI/ML workloads. Key services:")
bullet("Compute Engine: Virtual machines (like AWS EC2)")
bullet("Cloud Run: Run containers in a fully managed serverless environment — very developer-friendly")
bullet("Cloud Functions: Serverless functions (like AWS Lambda)")
bullet("Firebase: Real-time database, auth, hosting — popular for mobile apps")
bullet("BigQuery: Serverless data warehouse for analytics at petabyte scale")
bullet("Vertex AI: Machine learning platform — training, deployment, MLOps")
bullet("GKE (Google Kubernetes Engine): Managed Kubernetes — the best-in-class managed K8s")

h2("9.4 Microsoft Azure")
body("Azure is dominant in enterprise, especially Microsoft-stack companies:")
bullet("Azure App Service: Host web apps without managing VMs")
bullet("Azure Functions: Serverless compute")
bullet("Azure SQL Database: Managed SQL Server in the cloud")
bullet("Azure DevOps: Pipelines, repos, boards — very popular in enterprises")
bullet("Azure Active Directory: Enterprise identity management")
note("If you want to work in a large corporation or enterprise, Azure certifications (AZ-900 fundamentals) are valuable. Microsoft uses Azure itself and many enterprise companies follow.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — DOCKER & K8S
# ════════════════════════════════════════════════════════════════════════════
h1("10. Containers & Orchestration (Docker, Kubernetes)")

h2("10.1 Docker — Containerisation")
body("Docker is a platform for packaging applications and their dependencies into isolated, portable containers. A Docker container runs the same everywhere — your laptop, a CI server, or production.")
bullet("Problem it solves: 'It works on my machine' — containers eliminate environment differences")
bullet("Image: A blueprint/template for a container (like a class in OOP)")
bullet("Container: A running instance of an image (like an object/instance)")
bullet("Dockerfile: The recipe that defines how to build an image")
bullet("Docker Hub: Public registry for sharing Docker images")
bullet("docker-compose: Define and run multi-container apps (app + database + redis) locally")

h3("Essential Docker Commands")
make_table(
    ["Command", "What It Does"],
    [
        ["docker build -t myapp .", "Build an image from the current directory's Dockerfile"],
        ["docker run -p 3000:3000 myapp", "Run a container, mapping host port 3000 to container port 3000"],
        ["docker ps", "List running containers"],
        ["docker logs <container>", "See stdout/stderr from a running container"],
        ["docker exec -it <container> bash", "Get a shell inside a running container"],
        ["docker compose up -d", "Start all services defined in docker-compose.yml"],
        ["docker compose down", "Stop and remove all containers"],
        ["docker pull postgres:16", "Download an image from Docker Hub"],
    ],
    col_widths=[2.8, 3.7]
)
tip("Install Docker Desktop and replace your local database installations with Docker containers. Run postgres, redis, and mongodb via docker-compose.yml — this is how every professional development environment works.")

h2("10.2 Kubernetes (K8s) — Container Orchestration")
body("Kubernetes automates the deployment, scaling, and management of containerised applications across a cluster of machines. It is used when you have many containers and need them to be reliable, scalable, and self-healing.")
make_table(
    ["Concept", "What It Is"],
    [
        ["Pod", "The smallest unit — one or more containers that share network and storage"],
        ["Deployment", "Declares desired state: 'I want 3 replicas of this app running'"],
        ["Service", "Stable network endpoint to access a set of pods (load balances between them)"],
        ["Ingress", "HTTP/HTTPS routing rules — like Nginx but Kubernetes-native"],
        ["ConfigMap", "Store non-secret configuration (environment variables)"],
        ["Secret", "Store sensitive data (API keys, passwords) encrypted"],
        ["Namespace", "Virtual cluster within a cluster — separate environments (dev, staging, prod)"],
        ["HPA", "Horizontal Pod Autoscaler — automatically scales based on CPU/memory"],
        ["Helm", "Package manager for Kubernetes — install complex apps with one command"],
    ],
    col_widths=[1.8, 4.7]
)
note("Kubernetes is a complex system. You do not need deep K8s expertise to get a developer job — but understanding the concepts is asked in senior interviews. For now: learn Docker thoroughly, learn basic kubectl commands, and understand what Kubernetes solves.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 11 — SECURITY
# ════════════════════════════════════════════════════════════════════════════
h1("11. Security & Authentication")
body("Security is not optional. A single vulnerability can destroy a company's reputation and expose millions of users. Every professional developer must understand the most common attack vectors and how to prevent them.")

h2("11.1 The OWASP Top 10 — Must Know")
body("OWASP (Open Web Application Security Project) publishes the 10 most critical web application security risks. Employers and security auditors use this list:")
make_table(
    ["OWASP Risk", "What It Is", "How to Prevent"],
    [
        ["A01: Broken Access Control", "Users can access resources they shouldn't", "Check permissions server-side on every request, principle of least privilege"],
        ["A02: Cryptographic Failures", "Sensitive data exposed due to weak encryption", "Use TLS/HTTPS everywhere, bcrypt for passwords, never store plaintext"],
        ["A03: Injection (SQL, Command)", "Attacker injects malicious code into queries", "Use parameterised queries/ORM, never concatenate user input into SQL"],
        ["A04: Insecure Design", "Security not considered in architecture", "Threat model during design, security requirements from the start"],
        ["A05: Security Misconfiguration", "Default passwords, open ports, verbose errors", "Disable defaults, harden configs, use secrets managers"],
        ["A06: Vulnerable Components", "Using libraries with known CVEs", "Dependabot, npm audit, pip-audit, keep dependencies updated"],
        ["A07: Auth & Session Failures", "Weak passwords, no MFA, session fixation", "bcrypt, rate limiting, secure cookies, MFA"],
        ["A08: Software & Data Integrity", "Malicious code in CI/CD or dependencies", "Verify signatures, use trusted registries, SBOM"],
        ["A09: Logging Failures", "Not logging or monitoring attacks", "Log all auth attempts, anomalies, errors — never log sensitive data"],
        ["A10: Server-Side Request Forgery", "Server fetches URLs controlled by attacker", "Validate and allowlist URLs, block internal network requests"],
    ],
    col_widths=[1.8, 2.0, 2.7]
)

h2("11.2 Authentication vs Authorisation")
bullet("Authentication (AuthN): Who are you? — Verifying identity (login with email+password, OAuth, biometrics)")
bullet("Authorisation (AuthZ): What can you do? — Checking permissions (can this user delete this post?)")

h2("11.3 JWT — JSON Web Tokens")
body("JWTs are the most common mechanism for stateless authentication in APIs. Structure: Header.Payload.Signature")
bullet("Header: Algorithm and token type (e.g., HS256)")
bullet("Payload: Claims — userId, roles, expiry (exp), issued at (iat) — NOT for sensitive data")
bullet("Signature: Cryptographic signature — verifies the token wasn't tampered with")
bullet("Access Token: Short-lived (15 mins) — used to authenticate API requests")
bullet("Refresh Token: Long-lived (7-30 days) — used to get a new access token when it expires")
bullet("Store in: httpOnly cookies (not localStorage — XSS vulnerable)")
note("Common JWT mistakes: storing in localStorage, not validating expiry, using algorithm:none, not rotating refresh tokens on use. Know these for interviews.")

h2("11.4 OAuth 2.0 & OpenID Connect")
body("OAuth 2.0 is an authorisation framework. OpenID Connect (OIDC) adds authentication on top. Used for 'Sign in with Google/GitHub/etc.'")
make_table(
    ["Flow", "Use Case"],
    [
        ["Authorisation Code Flow", "Server-side web apps — the most secure flow"],
        ["PKCE (Proof Key for Code Exchange)", "Single-page apps and mobile — use instead of Implicit Flow"],
        ["Client Credentials Flow", "Machine-to-machine — backend service calling another service"],
        ["Refresh Token Flow", "Getting a new access token without re-login"],
    ],
    col_widths=[2.5, 4.0]
)

h2("11.5 Auth Libraries & Platforms")
make_table(
    ["Tool", "Type", "Best For"],
    [
        ["NextAuth.js (Auth.js)", "Library", "Next.js apps — easy social login, email magic links, adapters for Prisma/Drizzle"],
        ["Clerk", "SaaS", "Best DX, pre-built UI, organisations — best for freelance/SaaS"],
        ["Auth0", "SaaS", "Enterprise features, social login, MFA — very popular in companies"],
        ["Supabase Auth", "SaaS", "Comes with Supabase — easiest if using Supabase"],
        ["Firebase Auth", "SaaS", "Mobile-first, Google ecosystem"],
        ["Passport.js", "Library", "Node.js — 500+ strategy plugins for any provider"],
    ],
    col_widths=[1.7, 0.9, 4.1]
)
tip("For freelance SaaS products: Clerk is the fastest to set up and has the best user experience. For Next.js projects where you want full control: NextAuth.js (now Auth.js). Both are free to start.")

h2("11.6 Essential Security Practices")
bullet("HTTPS everywhere: Never serve HTTP in production — use Let's Encrypt (free) for SSL certificates")
bullet("Environment variables: Never hardcode API keys — use .env files locally, secrets manager in production")
bullet("Input validation: Validate and sanitise ALL user input on the server — never trust the frontend")
bullet("Rate limiting: Limit auth endpoints to 5-10 requests per minute per IP")
bullet("CORS: Only allow your frontend domain to call your API")
bullet("Security headers: Content-Security-Policy, X-Frame-Options, X-XSS-Protection, HSTS")
bullet("Dependency auditing: Run npm audit or pip-audit regularly — keep dependencies updated")
bullet("bcrypt for passwords: Never store plain or MD5-hashed passwords — use bcrypt with at least 10 rounds")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 12 — TESTING
# ════════════════════════════════════════════════════════════════════════════
h1("12. Testing — Unit, Integration & End-to-End")
body("Testing is the mark of a professional developer. Untested code is technical debt. Employers in serious companies will ask you about your testing experience, and many require it as part of the hiring process.")

h2("12.1 The Testing Pyramid")
body("The testing pyramid describes the ideal distribution of test types:")
bullet("Unit tests (base — most numerous): Test individual functions, classes, or components in isolation. Fast, cheap, many.")
bullet("Integration tests (middle): Test how multiple units work together (e.g., API endpoint + database query). Slower, fewer.")
bullet("End-to-End tests (top — fewest): Test complete user journeys through the UI. Slow, expensive, critical paths only.")
note("A common mistake is writing too many E2E tests. They are slow and brittle. Prefer many fast unit tests and targeted E2E tests for critical flows.")

h2("12.2 Frontend Testing Stack")
make_table(
    ["Tool", "Type", "What It Tests", "Notes"],
    [
        ["Vitest", "Unit/Integration", "JavaScript functions, hooks, utilities", "Fast, Vite-native — REPLACE Jest for new projects"],
        ["Jest", "Unit/Integration", "JavaScript functions, hooks, utilities", "Older standard — still widely used"],
        ["React Testing Library", "Component", "Component rendering and user interactions", "Tests behaviour, not implementation — MUST LEARN"],
        ["Playwright", "E2E", "Full browser automation — Chrome, Firefox, Safari", "Best E2E tool today — RECOMMENDED"],
        ["Cypress", "E2E", "Browser automation with good DX", "Popular, but Playwright is faster and more capable"],
        ["Storybook", "Visual", "UI component development and visual testing", "Industry standard for component libraries"],
        ["MSW (Mock Service Worker)", "API Mocking", "Mock API calls in tests and development", "Essential for testing components that fetch data"],
    ],
    col_widths=[1.6, 1.2, 2.2, 2.5]
)

h2("12.3 Backend Testing Stack (Python)")
make_table(
    ["Tool", "Type", "Notes"],
    [
        ["pytest", "Unit/Integration", "The Python testing standard — fixtures, parametrize, plugins are excellent"],
        ["httpx / requests", "HTTP client for tests", "Call your FastAPI/Flask endpoints in tests"],
        ["pytest-asyncio", "Async tests", "Test async FastAPI endpoints"],
        ["factory_boy", "Test data factories", "Generate test data without manual fixtures"],
        ["Faker", "Fake data", "Generate realistic fake names, emails, addresses for tests"],
        ["coverage.py", "Test coverage", "Measure what percentage of code is tested — aim for 70-80%+"],
    ],
    col_widths=[1.8, 1.4, 3.3]
)

h2("12.4 Backend Testing Stack (Node.js)")
make_table(
    ["Tool", "Notes"],
    [
        ["Vitest", "Unit tests — fast, modern"],
        ["Supertest", "HTTP integration tests — call your Express/Fastify endpoints in tests"],
        ["Nock", "Mock HTTP requests to third-party APIs in tests"],
    ],
    col_widths=[1.8, 4.7]
)
tip("Start here: Write tests for a project you have already built. Aim for 50% coverage. The process of writing tests will reveal bugs and make you understand your own code much better.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 13 — SYSTEM DESIGN
# ════════════════════════════════════════════════════════════════════════════
h1("13. System Design & Architecture")
body("System design is how senior developers think about building software at scale. You will face system design questions in interviews — even as a junior/mid developer. Understanding these concepts makes you a much better developer at any level.")

h2("13.1 Fundamental Concepts")
make_table(
    ["Concept", "Explanation"],
    [
        ["Scalability", "Ability to handle more load. Horizontal scaling = more servers. Vertical scaling = bigger server."],
        ["Availability", "Percentage of time a system is operational. 99.9% = 8.7 hrs downtime/year. 99.99% = 52 mins/year."],
        ["Reliability", "Consistency of correct operation. A system can be available but return wrong data."],
        ["Latency", "Time for a request to complete. P99 latency = 99th percentile — worst-case for most users."],
        ["Throughput", "Number of requests a system can handle per second (RPS/QPS)."],
        ["CAP Theorem", "In distributed systems, you can only guarantee 2 of: Consistency, Availability, Partition tolerance."],
        ["Single Point of Failure", "Any component whose failure brings down the system — avoid with redundancy."],
        ["Load Balancing", "Distribute traffic across multiple servers. Round-robin, least-connections, consistent hashing."],
        ["Caching", "Store computed results to avoid repeating expensive operations. Cache invalidation is hard."],
        ["CDN", "Content Delivery Network — serve static assets from edge locations near users."],
        ["Database Replication", "Copy data to multiple servers. Read replicas improve read throughput."],
        ["Database Sharding", "Split data across multiple database instances by key (e.g., userId % n)."],
        ["Message Queue", "Async communication between services. Decouple producers from consumers."],
        ["Rate Limiting", "Control how many requests a client can make in a time window."],
        ["Circuit Breaker", "Stop calling a failing service to prevent cascade failures."],
    ],
    col_widths=[1.8, 4.7]
)

h2("13.2 Architecture Patterns")
make_table(
    ["Pattern", "Description", "Pros", "Cons"],
    [
        ["Monolith", "Single deployable unit with all code", "Simple, easy to develop and debug", "Hard to scale, deploy, or work on in large teams"],
        ["Microservices", "Independent services per domain", "Scale independently, tech diversity, team autonomy", "Distributed complexity, network latency, ops overhead"],
        ["Modular Monolith", "One app, but well-separated modules", "Best of both worlds for small-medium teams", "Requires discipline to maintain module boundaries"],
        ["Serverless", "Functions run on demand", "Auto-scaling, pay-per-use, no server management", "Cold starts, vendor lock-in, debugging harder"],
        ["Event-Driven", "Services communicate via events/messages", "Loose coupling, async processing, audit trails", "Eventual consistency, harder to debug"],
        ["Hexagonal (Ports & Adapters)", "Domain logic isolated from infrastructure", "Testable, flexible — swap database/framework easily", "More boilerplate"],
    ],
    col_widths=[1.5, 1.8, 1.8, 1.9]
)
tip("For your level: build modular monoliths. Structure your code so modules could be extracted into microservices if needed, but keep it as one deployable application. This is what most startups do right.")

h2("13.3 How to Answer System Design Interviews")
body("The standard approach to a system design question (e.g., 'Design a URL shortener like bit.ly'):")
bullet("Step 1 — Clarify requirements (5 mins): Functional requirements (what does it do?), Non-functional (scale, availability, latency?)")
bullet("Step 2 — Estimate scale (2 mins): Users, requests per second, storage needed — back-of-envelope math")
bullet("Step 3 — API design (5 mins): What endpoints/methods will the system expose?")
bullet("Step 4 — Database schema (5 mins): What data do we store and how?")
bullet("Step 5 — High-level design (10 mins): Draw the main components — clients, servers, databases, caches")
bullet("Step 6 — Deep dive (10 mins): Interviewers pick a component to discuss in detail")
bullet("Step 7 — Identify bottlenecks (5 mins): What breaks first at scale? How would you fix it?")
note("Practice on: Grokking the System Design Interview (paid, worth it), ByteByteGo (Alex Xu's free YouTube content), and systemdesign.one.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 14 — CS FUNDAMENTALS
# ════════════════════════════════════════════════════════════════════════════
h1("14. Computer Science Fundamentals")
body("As a CS graduate, you have studied these academically. The challenge is applying them in coding interviews and knowing them well enough to discuss in system design contexts.")

h2("14.1 Data Structures — Interview Priority")
make_table(
    ["Data Structure", "Key Operations", "When to Use", "Frequency in Interviews"],
    [
        ["Array / List", "Access O(1), Search O(n), Insert O(n)", "Default — ordered data", "★★★★★"],
        ["Hash Map / Dict", "Get/Set O(1) average", "O(1) lookups, counting, grouping", "★★★★★"],
        ["Hash Set", "Contains O(1) average", "Uniqueness checks, visited tracking", "★★★★★"],
        ["Stack", "Push/Pop O(1)", "DFS, undo, bracket matching", "★★★★☆"],
        ["Queue / Deque", "Enqueue/Dequeue O(1)", "BFS, sliding window", "★★★★☆"],
        ["Linked List", "Insert O(1), Access O(n)", "Rarely best choice, but asked often", "★★★☆☆"],
        ["Binary Tree", "Traversal O(n)", "Hierarchical data", "★★★★☆"],
        ["Binary Search Tree", "Search O(log n) avg", "Sorted data with frequent insert/delete", "★★★☆☆"],
        ["Heap / Priority Queue", "Insert O(log n), Min/Max O(1)", "Top K problems, scheduling", "★★★★☆"],
        ["Trie", "Insert/Search O(m)", "Autocomplete, prefix search", "★★★☆☆"],
        ["Graph", "Varies by representation", "Relationships, paths, networks", "★★★★☆"],
    ],
    col_widths=[1.6, 1.8, 2.0, 1.1]
)

h2("14.2 Algorithms — Essential Patterns")
make_table(
    ["Pattern", "Description", "Example Problems"],
    [
        ["Two Pointers", "Two pointers moving toward each other or at different speeds", "Two Sum (sorted), Remove Duplicates, Container With Most Water"],
        ["Sliding Window", "Fixed or variable window over an array/string", "Maximum Subarray, Longest Substring Without Repeating"],
        ["Binary Search", "Search sorted data in O(log n)", "Search in Rotated Array, Find Peak Element"],
        ["BFS (Breadth-First)", "Level-order traversal using queue", "Shortest Path, Level Order Traversal, Word Ladder"],
        ["DFS (Depth-First)", "Explore all paths using stack or recursion", "All Paths, Number of Islands, Clone Graph"],
        ["Dynamic Programming", "Break into subproblems, cache results", "Fibonacci, Coin Change, Longest Common Subsequence"],
        ["Backtracking", "Explore all options, prune invalid paths", "N-Queens, Sudoku, Permutations, Subsets"],
        ["Greedy", "Always pick locally optimal choice", "Activity Selection, Huffman Encoding, Jump Game"],
        ["Merge Intervals", "Combine overlapping intervals", "Merge Intervals, Insert Interval, Meeting Rooms"],
        ["Topological Sort", "Order nodes with dependencies", "Course Schedule, Build Order, Task Dependency"],
    ],
    col_widths=[1.8, 2.2, 2.5]
)
tip("Do at least 100 LeetCode problems before serious job applications. Focus on Easy and Medium. Do these in Python (your strongest language). NeetCode.io has a free structured roadmap — use it.")

h2("14.3 Big-O Complexity — Quick Reference")
make_table(
    ["Notation", "Name", "Example"],
    [
        ["O(1)", "Constant", "Hash map lookup, array index access"],
        ["O(log n)", "Logarithmic", "Binary search, balanced BST operations"],
        ["O(n)", "Linear", "Linear search, single loop through array"],
        ["O(n log n)", "Linearithmic", "Efficient sorting (merge sort, heap sort)"],
        ["O(n²)", "Quadratic", "Nested loops, bubble sort"],
        ["O(2ⁿ)", "Exponential", "Recursive Fibonacci (naive), subset generation"],
        ["O(n!)", "Factorial", "Permutations of n elements"],
    ],
    col_widths=[1.2, 1.5, 4.0]
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 15 — PORTFOLIO & JOB STRATEGY
# ════════════════════════════════════════════════════════════════════════════
h1("15. Soft Skills, Portfolio & Job Strategy")

h2("15.1 Building a Portfolio That Gets Interviews")
body("Your portfolio is the most important factor for getting your first professional role. It proves you can build things.")

h3("What to Build (Priority Order)")
make_table(
    ["Project Type", "What to Include", "Why It Impresses"],
    [
        ["Full-Stack SaaS Clone", "A simplified clone of a real product (Trello, Notion, Twitter/X)", "Shows you understand real-world features: auth, CRUD, real-time, subscriptions"],
        ["AI-Powered App", "App using OpenAI/Claude API: writing tool, chat, code reviewer", "AI is the hottest skill right now — shows you understand LLMs"],
        ["E-Commerce Store", "Products, cart, Stripe checkout, orders, admin panel", "Extremely common client request — proves you can handle payments"],
        ["REST API with Docs", "A well-documented FastAPI/Express API with Swagger/OpenAPI", "Shows backend depth — many devs skip documentation"],
        ["Open Source Contribution", "PR to a popular open source project", "Shows collaboration, reading others' code, professional Git workflow"],
    ],
    col_widths=[1.7, 2.5, 2.3]
)

h3("Portfolio Website Must-Haves")
bullet("About section: Who you are, your tech stack, what types of projects you build")
bullet("Projects: 3-6 projects with live demo links AND GitHub repo links — no repos with no README")
bullet("Tech stack shown on each project: Don't make them guess what you used")
bullet("Contact form or email: Make it easy to reach you")
bullet("Blog (optional but powerful): Writing about what you learn shows communication skills")

h3("README Files — Your Silent Sales Pitch")
body("Every project on GitHub needs a professional README with:")
bullet("What the project does (one sentence)")
bullet("Tech stack badges")
bullet("Screenshots or a GIF of it working")
bullet("Installation and setup instructions")
bullet("Features list")
bullet("Link to live demo")
tip("Use shields.io for badges and GitHub README templates. A professional README can triple the number of people who actually look at your project.")

h2("15.2 CV / Resume Best Practices")
make_table(
    ["Section", "What to Include", "Common Mistakes"],
    [
        ["Header", "Name, email, GitHub, LinkedIn, portfolio URL", "Including photo (in some countries), full address, date of birth"],
        ["Summary", "2-3 sentences: your stack, experience level, what you're looking for", "Too generic: 'passionate developer looking for opportunities'"],
        ["Skills", "Group by category: Languages, Frameworks, Databases, Cloud, Tools", "Listing 50 skills you barely know — stick to what you can interview on"],
        ["Projects", "Project name, tech stack, 2-3 bullet points on what YOU did and metrics", "Just listing the project name with no context"],
        ["Experience", "Internship: company, role, dates, 3-5 bullet points with impact metrics", "Only job duties — include what you achieved"],
        ["Education", "Degree, institution, graduation year — no grades unless excellent", "Listing every course you took"],
    ],
    col_widths=[1.3, 2.5, 2.7]
)
note("Use the XYZ formula for bullet points: Accomplished [X] as measured by [Y], by doing [Z]. Example: 'Reduced page load time by 40% by implementing code splitting and image optimisation, improving Lighthouse score from 60 to 95.'")

h2("15.3 The Job Search Strategy")
h3("Where to Apply")
bullet("LinkedIn: Most professional roles, company pages, Easy Apply — volume play")
bullet("Indeed: Wide variety of roles, including small companies not on LinkedIn")
bullet("Glassdoor: Apply and also read company reviews before interviewing")
bullet("AngelList (Wellfound): Startups — often more flexible on experience requirements")
bullet("Stack Overflow Jobs: Developer-specific, often quality roles")
bullet("Remote.co / We Work Remotely / Remote OK: Remote roles globally")
bullet("Local tech meetups: In-person networking — underused but effective")
bullet("Cold outreach: Find developers at companies you like on LinkedIn and message them")

h3("The Application Strategy for Getting Your First Role")
bullet("Apply to 10-15 jobs per week — not 2-3. Volume matters when you are new.")
bullet("Tailor your CV for each role — match their keywords exactly.")
bullet("Apply within 24 hours of a posting — early applicants have higher success rates.")
bullet("Follow up: If no response after 1 week, a polite follow-up email is professional, not pushy.")
bullet("Track everything: Use a spreadsheet — company, role, date applied, status, next step.")

h3("Interview Process — What to Expect")
make_table(
    ["Stage", "What Happens", "How to Prepare"],
    [
        ["Application screening", "HR reviews CV keywords", "Match job description keywords in your CV"],
        ["HR phone screen", "15-30 min culture fit, availability, salary", "Research company, know your numbers"],
        ["Technical assessment", "Take-home project or LeetCode-style test", "Practice on LeetCode, build a demo project quickly"],
        ["Technical interview", "Live coding + system design + technical Q&A", "LeetCode 150, system design study, CS fundamentals"],
        ["Behavioural interview", "STAR method questions about past experience", "Prepare 5-7 stories using STAR format"],
        ["Final round", "Meet the team, culture fit, negotiate offer", "Research team, prepare salary negotiation"],
    ],
    col_widths=[1.6, 2.0, 3.0]
)
tip("The STAR method for behavioural questions: Situation (context), Task (your role), Action (what you did), Result (the outcome with metrics). Prepare stories for: teamwork, failure, leadership, conflict, technical challenge.")

h2("15.4 Salary Negotiation")
body("Salary negotiation is expected and professional. Developers who negotiate earn significantly more over their career:")
bullet("Always negotiate: Employers expect it. The worst they can say is no.")
bullet("Let them make the first offer: If asked your expectations, say 'I'm flexible — what is the range for this role?'")
bullet("Research market rate: levels.fyi, Glassdoor, LinkedIn Salary, local developer survey data")
bullet("Negotiate the full package: base salary, equity, signing bonus, remote flexibility, learning budget")
bullet("Counter with a specific number: '€55,000' beats 'something higher' — be specific")
bullet("Get it in writing: Always confirm verbal offers via email before resigning from current job")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 16 — FREELANCING
# ════════════════════════════════════════════════════════════════════════════
h1("16. Freelancing & Self-Employment Guide")
body("Freelancing gives you freedom, higher hourly rates, and the ability to choose your clients. But it requires business skills alongside technical ones. This section covers everything to build a sustainable freelance developer practice.")

h2("16.1 How to Start Freelancing (Even as a Beginner)")
bullet("Build 2-3 portfolio projects first: Clients need proof you can build things. The portfolio projects from Section 15 apply here.")
bullet("Start with your network: Your first clients are usually people you already know — friends, family, former classmates, LinkedIn connections.")
bullet("Take your first project at a discounted rate: Getting a case study with real results matters more than your first rate.")
bullet("Document everything you build: Write case studies, take screenshots, ask for testimonials.")
bullet("Specialise early: 'React developer for SaaS companies' is easier to sell than 'full-stack developer'. Niche down.")

h2("16.2 Where to Find Freelance Clients")
make_table(
    ["Platform/Channel", "Type", "Best For", "Competition"],
    [
        ["Upwork", "Marketplace", "Long-term contracts, international clients", "Very high — hard to start"],
        ["Toptal", "Curated marketplace", "Top 3% developers, high rates", "Extreme vetting process"],
        ["Freelancer.com", "Marketplace", "Short projects, beginners", "Very high, race to bottom pricing"],
        ["Fiverr", "Marketplace", "Productised services, US clients", "High, but niches work well"],
        ["LinkedIn", "Social/Outreach", "B2B, agency work, referrals", "Low if you approach correctly"],
        ["Cold email/DM", "Outreach", "Targeted ideal clients", "No competition — you pick them"],
        ["Local businesses", "In-person", "Small business websites, systems", "Low — most ignore them"],
        ["Developer communities", "Network", "Overflow work from other devs, referrals", "Low — relationships-based"],
        ["Twitter/X & Reddit", "Content/Community", "Inbound leads from showing your work", "Low — value-based"],
        ["Your own website + SEO", "Inbound", "Long-term lead generation", "None — passive"],
    ],
    col_widths=[1.7, 1.2, 2.1, 1.5]
)
tip("The fastest path to freelance clients is warm outreach to your network + LinkedIn DMs to small business owners who have an outdated website or no web presence. One message: 'I noticed your website is X — I can modernise it. Here are 3 sites I built recently.'")

h2("16.3 Pricing Your Services")
make_table(
    ["Model", "How It Works", "Pros", "Cons"],
    [
        ["Hourly Rate", "Bill per hour worked", "Simple, protected against scope creep", "Penalises efficiency, awkward for clients"],
        ["Fixed Project Price", "Quote a lump sum for a defined scope", "Predictable for client, reward for efficiency", "Risk of scope creep — use contracts"],
        ["Retainer", "Monthly fee for ongoing work/availability", "Stable income, relationship building", "Requires clear deliverables per month"],
        ["Value-Based Pricing", "Price based on value delivered to client", "Highest earning potential", "Requires understanding client ROI"],
    ],
    col_widths=[1.5, 2.0, 1.8, 2.2]
)

h3("Rate Benchmarks (Europe/Remote, 2024-2025)")
make_table(
    ["Experience Level", "Hourly Rate (EUR)", "Monthly (Full-Time Equivalent)"],
    [
        ["Beginner (0-1 years)", "€25–45/hr", "€4,000–7,200/month"],
        ["Junior-Mid (1-3 years)", "€45–80/hr", "€7,200–12,800/month"],
        ["Mid-Senior (3-6 years)", "€80–130/hr", "€12,800–20,800/month"],
        ["Senior (6+ years)", "€130–200+/hr", "€20,800–32,000+/month"],
    ],
    col_widths=[2.2, 1.8, 2.5]
)
note("These are freelance rates, not salaried rates. They are higher to compensate for taxes, self-employment costs, non-billable time, and no employer benefits. Factor in ~30% for taxes and 20% for non-billable time.")

h2("16.4 The Freelance Business Foundation")
h3("Legal & Financial (Germany/EU Context)")
bullet("Business registration: In Germany, register as Freiberufler (freelancer) with the Finanzamt — simpler than GmbH for software developers")
bullet("VAT/Umsatzsteuer: Register for Umsatzsteuer if revenue exceeds €22,000/year. EU clients: check reverse charge rules.")
bullet("Kleinunternehmerregelung: Under €22,000/year, you can opt out of charging VAT — simpler for starting out")
bullet("Separate business bank account: Keep business finances completely separate from personal")
bullet("Accounting: Use DATEV, Lexoffice, or SevDesk for German freelancers")
bullet("Tax professional: A Steuerberater specialising in freelancers is worth the cost — they save you more than they charge")
bullet("Insurance: Berufshaftpflichtversicherung (professional liability insurance) is essential for client contracts")

h3("Contracts — Never Work Without One")
bullet("Always use a written contract for any project over €500")
bullet("Define scope exactly: What is included, what is not")
bullet("Payment terms: Milestone payments (30% upfront, 40% mid-project, 30% on delivery) protect you")
bullet("Intellectual property: Client owns the code on full payment — specify this explicitly")
bullet("Change request process: How scope changes are priced and approved")
bullet("Termination clause: What happens if the project is cancelled")
tip("Use HelloSign, DocuSign, or PandaDoc for digital signatures. There are free contract templates for software freelancers at freelancermap.com and Bonsai.")

h2("16.5 Productised Services — Scale Without Burning Out")
body("Productised services are freelance offerings packaged as fixed-price products with defined scope and delivery time. They are easier to sell and deliver than custom projects:")
make_table(
    ["Service", "Price Range", "Delivery Time"],
    [
        ["Landing page (React + Tailwind + animations)", "€800–2,500", "3–5 days"],
        ["Full business website (5-10 pages, CMS)", "€2,500–6,000", "1–2 weeks"],
        ["SaaS MVP (auth + core feature + payments)", "€5,000–15,000", "4–8 weeks"],
        ["API integration (Stripe, third-party service)", "€500–2,000", "1–3 days"],
        ["Performance audit + optimisation", "€800–2,500", "3–5 days"],
        ["Monthly website maintenance retainer", "€300–800/month", "Ongoing"],
    ],
    col_widths=[2.5, 1.5, 1.5]
)

h2("16.6 Technology Stack for Freelance Products")
body("For maximum productivity on freelance projects, standardise on a small, well-chosen stack:")
bullet("Frontend: Next.js + TypeScript + Tailwind CSS + shadcn/ui")
bullet("Auth: Clerk (fastest DX) or NextAuth.js")
bullet("Database: Supabase (Postgres + storage + auth) or PlanetScale/Neon")
bullet("ORM: Prisma")
bullet("Payments: Stripe")
bullet("Email: Resend (easiest API) or SendGrid")
bullet("Hosting: Vercel (frontend/serverless) + Railway (if you need a dedicated backend)")
bullet("CMS (content management): Sanity.io, Contentful, or Payload CMS")
bullet("Forms: React Hook Form + Zod")
bullet("Analytics: Plausible (privacy-first, you already use Umami — similar philosophy)")
note("Mastering ONE stack and delivering quality products fast is worth more than knowing 10 stacks superficially. The above stack lets you build full-featured products in days, not months.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 17 — CERTIFICATIONS
# ════════════════════════════════════════════════════════════════════════════
h1("17. Free Certifications Master List")
body("Certifications alone don't get you a job, but they signal commitment, structured learning, and cover specific domains employers value. All certifications below have zero cost or free preparation resources.")

h2("17.1 Cloud & Infrastructure")
make_table(
    ["Certification", "Provider", "Cost", "Duration", "Value"],
    [
        ["AWS Cloud Practitioner (CLF-C02)", "AWS", "Exam $100 — free prep via AWS Skill Builder", "2-4 weeks study", "★★★★★ — Best starter cloud cert"],
        ["Google Cloud Digital Leader", "Google", "Exam $99 — free prep via Google Cloud Skills Boost", "2-3 weeks", "★★★★☆ — Good alternative to AWS"],
        ["Microsoft AZ-900 Azure Fundamentals", "Microsoft", "Exam $165 — free prep via Microsoft Learn", "2-3 weeks", "★★★★☆ — Enterprise value"],
        ["Google Associate Cloud Engineer", "Google", "Exam $200 — free prep available", "6-8 weeks", "★★★★★ — Strong practical cert"],
        ["AWS Solutions Architect Associate", "AWS", "Exam $150 — free prep via FreeCodeCamp YouTube", "6-10 weeks", "★★★★★ — Most respected AWS cert"],
        ["Kubernetes & Cloud Native Associate (KCNA)", "CNCF", "Exam $250 — free prep via CNCF", "4-6 weeks", "★★★★☆ — Growing demand"],
    ],
    col_widths=[2.5, 1.2, 1.8, 1.3, 0.9]
)

h2("17.2 100% Free Certificates (No Exam Fee)")
make_table(
    ["Certificate", "Provider", "Platform", "Topics Covered"],
    [
        ["Responsive Web Design", "freeCodeCamp", "freecodecamp.org", "HTML, CSS, Flexbox, Grid, Accessibility"],
        ["JavaScript Algorithms & Data Structures", "freeCodeCamp", "freecodecamp.org", "JS, ES6, algorithms, regex, OOP"],
        ["Front End Development Libraries", "freeCodeCamp", "freecodecamp.org", "React, Redux, Bootstrap, jQuery"],
        ["Back End Development & APIs", "freeCodeCamp", "freecodecamp.org", "Node.js, Express, MongoDB, REST APIs"],
        ["Relational Database", "freeCodeCamp", "freecodecamp.org", "SQL, PostgreSQL, Bash, Git"],
        ["Quality Assurance", "freeCodeCamp", "freecodecamp.org", "Chai, Mocha, testing"],
        ["Scientific Computing with Python", "freeCodeCamp", "freecodecamp.org", "Python fundamentals, data structures"],
        ["Python for Everybody", "University of Michigan", "Coursera (audit free)", "Python basics, web scraping, databases"],
        ["Django for Everybody", "University of Michigan", "Coursera (audit free)", "Django, HTML, SQL, JavaScript"],
        ["Full Stack Open", "University of Helsinki", "fullstackopen.com", "React, Node, GraphQL, TypeScript, Testing"],
        ["CS50x", "Harvard", "edx.org (audit free)", "Computer science foundations"],
        ["CS50 Web Programming", "Harvard", "edx.org (audit free)", "HTML, CSS, JS, SQL, Python, Django"],
        ["The Odin Project — Full Stack JavaScript", "The Odin Project", "theodinproject.com", "HTML, CSS, JS, Node, React, Databases"],
        ["Meta Frontend Developer Certificate", "Meta", "Coursera (audit free)", "React, Redux, UX, version control"],
        ["Meta Backend Developer Certificate", "Meta", "Coursera (audit free)", "Python, Django, APIs, databases"],
        ["IBM Full Stack Software Developer", "IBM", "Coursera (audit free)", "HTML, CSS, JS, React, Node, Python, Docker"],
        ["GitHub Foundations", "GitHub", "learn.microsoft.com", "Git, GitHub, collaboration, CI/CD"],
        ["Google UX Design Certificate", "Google", "Coursera (audit free)", "UX research, wireframing, Figma"],
        ["Linux Essentials", "Linux Professional Institute", "lpi.org (free materials)", "Linux fundamentals"],
        ["Cybersecurity Fundamentals", "IBM", "Coursera (audit free)", "Security basics, encryption, threats"],
    ],
    col_widths=[2.2, 1.8, 1.7, 2.0]
)

h2("17.3 Platform-Specific Free Badges & Certificates")
make_table(
    ["Platform", "What They Offer", "URL"],
    [
        ["AWS Skill Builder", "Free digital badges for cloud skills, free exam prep courses", "skillbuilder.aws"],
        ["Google Cloud Skills Boost", "Hands-on labs, skill badges, learning paths", "cloudskillsboost.google"],
        ["Microsoft Learn", "Badges, learning paths, exam sandboxes for Azure", "learn.microsoft.com"],
        ["MongoDB University", "Free MongoDB Developer, DBA, Data Analyst certificates", "learn.mongodb.com"],
        ["Redis University", "Free Redis certificates (Redis for JavaScript, Python, etc.)", "university.redis.io"],
        ["Scrum.org", "Free Scrum Guide, paid PSM I cert (~$150 — worth it)", "scrum.org"],
        ["HackerRank", "Free skill certificates in Python, SQL, Problem Solving, React", "hackerrank.com"],
        ["Kaggle", "Free certificates in Python, SQL, ML, Data Viz", "kaggle.com/learn"],
        ["Postman", "Free API testing and design certificates", "academy.postman.com"],
        ["Docker", "Free Docker Foundations certification", "docker.com/learning-path"],
    ],
    col_widths=[1.8, 3.0, 2.0]
)

h2("17.4 Recommended Certificate Priority Order")
body("Given your goals (job now, freelance later), collect these in this order:")
bullet("Month 1-2: freeCodeCamp JavaScript + React certificates (show to employers)")
bullet("Month 2-3: Full Stack Open certificate (highly respected in European tech companies)")
bullet("Month 3-4: AWS Cloud Practitioner (demonstrates cloud knowledge to employers)")
bullet("Month 4-5: HackerRank Python + Problem Solving certificates")
bullet("Month 5-6: Meta Backend Developer or IBM Full Stack (strong brand recognition)")
bullet("Month 6+: AWS Solutions Architect Associate (for senior/higher paying roles)")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 18 — STUDY PLAN
# ════════════════════════════════════════════════════════════════════════════
h1("18. Weekly Study Plan (30+ hrs/week)")
body("This is a structured 6-month plan at 30-35 hours per week. It prioritises getting you job-ready first, then freelance-ready second. Each phase builds on the previous.")

h2("Phase 1 — Weeks 1-4: Strengthen Your Core (Foundation)")
make_table(
    ["Day", "Morning (4 hrs)", "Afternoon (3 hrs)", "Evening (1 hr)"],
    [
        ["Mon", "Next.js App Router: official tutorial + build pages", "TanStack Query + Zustand: build a data fetching project", "LeetCode 1 easy problem"],
        ["Tue", "FastAPI: Build a REST API (CRUD for a resource)", "PostgreSQL: install, connect, practice SQL queries", "Read OWASP Top 10 article"],
        ["Wed", "React Hook Form + Zod: build a validated form", "SQLAlchemy + Alembic: ORM models + migration", "LeetCode 1 easy problem"],
        ["Thu", "Next.js API Routes + Prisma + Postgres full-stack CRUD", "Supabase: set up free project, connect to Next.js", "Watch system design YouTube (ByteByteGo)"],
        ["Fri", "Auth: Implement NextAuth.js (Google OAuth)", "Docker: Write Dockerfile for your app, docker-compose", "LeetCode 1 easy problem"],
        ["Sat", "Build Project: Full-stack app with auth + CRUD + Prisma", "Same project: deploy to Vercel + Railway", "Write GitHub README"],
        ["Sun", "Review week: document what you learned, fix stuck areas", "Portfolio: add completed project with case study", "Rest / light reading"],
    ],
    col_widths=[0.6, 2.5, 2.5, 1.4]
)
note("Phase 1 Goal: A deployed full-stack Next.js app with auth, database (Prisma + Postgres), deployed on Vercel + Railway. This alone is better than most graduate portfolios.")

h2("Phase 2 — Weeks 5-8: Deepen Backend + Testing")
make_table(
    ["Focus", "Topics", "Deliverable"],
    [
        ["Week 5", "FastAPI advanced: background tasks, middleware, file uploads, async", "FastAPI project with file upload, Celery background job"],
        ["Week 6", "Testing: pytest for backend, React Testing Library for frontend, Playwright E2E", "Test suite for your Phase 1 project (>60% coverage)"],
        ["Week 7", "Redis: caching, sessions, queues. BullMQ (Node) or Celery (Python)", "Add Redis caching to your API, implement a job queue"],
        ["Week 8", "GitHub Actions CI/CD: lint, test, deploy pipeline. Security: JWT, CORS, rate limiting", "CI/CD pipeline for your main project"],
    ],
    col_widths=[0.9, 2.8, 2.8]
)

h2("Phase 3 — Weeks 9-12: Cloud + DevOps + Portfolio Polish")
make_table(
    ["Focus", "Topics", "Deliverable"],
    [
        ["Week 9", "AWS fundamentals: S3, Lambda, RDS (free tier). AWS Cloud Practitioner prep", "Deploy static site to S3 + CloudFront. Start AWS cert prep"],
        ["Week 10", "Docker + deployment: containerise apps, deploy to Fly.io or Render with Docker", "Dockerised full-stack app deployed with docker-compose"],
        ["Week 11", "System design study: ByteByteGo, Grokking. Practice 2 design problems", "Design doc for a URL shortener + design doc for a chat app"],
        ["Week 12", "AWS Cloud Practitioner exam preparation + take the exam", "AWS Cloud Practitioner certificate"],
    ],
    col_widths=[0.9, 2.8, 2.8]
)

h2("Phase 4 — Weeks 13-16: Job Applications + Advanced Projects")
make_table(
    ["Focus", "Topics", "Deliverable"],
    [
        ["Week 13", "Build AI-powered project (OpenAI/Claude API integration in Next.js)", "Deployed AI app — writing tool, code reviewer, or chatbot"],
        ["Week 14", "LeetCode ramp-up: 15 Medium problems. Interview prep: STAR method stories", "Track record of 50 total LeetCode problems"],
        ["Week 15", "Job application sprint: apply to 10-15 roles. LinkedIn profile optimisation", "Applications sent, LinkedIn profile at 'All-Star' level"],
        ["Week 16", "Interview practice: mock interviews (Pramp, interviewing.io), code review skills", "2+ mock interviews done. Reviewed 3 open source PRs"],
    ],
    col_widths=[0.9, 2.8, 2.8]
)

h2("Phase 5 — Months 5-6: First Freelance Clients")
make_table(
    ["Focus", "Topics", "Deliverable"],
    [
        ["Month 5", "Stripe integration: build a subscription SaaS with Clerk + Stripe + Prisma", "Working SaaS template with auth + subscriptions"],
        ["Month 5", "Productised service packaging: write 3 service packages with clear pricing", "Public-facing services page on portfolio site"],
        ["Month 6", "Client outreach: 20 personalised messages to potential clients", "First paid project (even if small)"],
        ["Month 6", "Business setup: register as freelancer, open business account, get contract template", "Legal freelance structure ready to invoice"],
    ],
    col_widths=[0.9, 2.8, 2.8]
)

h2("18.1 Daily LeetCode Schedule")
body("LeetCode practice should be a daily habit — 1 problem/day minimum, 2-3 on good days:")
bullet("Weeks 1-6: All easy problems (Arrays, Strings, Hash Maps)")
bullet("Weeks 7-12: Medium problems by pattern (Two Pointers, Sliding Window, BFS, DFS)")
bullet("Weeks 13-20: Mix of medium + revisit missed problems")
bullet("Before interviews: Company-specific question lists on LeetCode")
tip("Use NeetCode.io — it organises LeetCode problems by pattern with free video explanations. Do the NeetCode 150 list.")

h2("18.2 Resources Directory")
make_table(
    ["Resource", "Type", "URL", "Cost"],
    [
        ["Full Stack Open", "Course", "fullstackopen.com", "Free"],
        ["The Odin Project", "Course", "theodinproject.com", "Free"],
        ["freeCodeCamp", "Course + Certs", "freecodecamp.org", "Free"],
        ["NeetCode.io", "LeetCode prep", "neetcode.io", "Free + paid"],
        ["ByteByteGo", "System Design", "blog.bytebytego.com + YouTube", "Free (YouTube)"],
        ["Roadmap.sh", "Learning roadmaps", "roadmap.sh", "Free"],
        ["MDN Web Docs", "Reference", "developer.mozilla.org", "Free"],
        ["FastAPI docs", "Reference", "fastapi.tiangolo.com", "Free"],
        ["Next.js docs", "Reference", "nextjs.org/docs", "Free"],
        ["PostgreSQL Tutorial", "Tutorial", "postgresqltutorial.com", "Free"],
        ["Redis University", "Certificates", "university.redis.io", "Free"],
        ["Docker Getting Started", "Tutorial", "docs.docker.com/get-started", "Free"],
        ["Stripe Docs", "Reference", "stripe.com/docs", "Free"],
        ["AWS Skill Builder", "Cloud training", "skillbuilder.aws", "Free tier"],
        ["Prisma Docs", "Reference", "prisma.io/docs", "Free"],
        ["Supabase Docs", "Reference", "supabase.com/docs", "Free"],
        ["JavaScript.info", "JS Deep Dive", "javascript.info", "Free"],
        ["TypeScript Handbook", "Reference", "typescriptlang.org/docs", "Free"],
    ],
    col_widths=[1.8, 1.0, 2.5, 1.2]
)

page_break()

# Final page
h1("Final Words")
body("You are starting from a better position than you think. You already have a CS degree, real internship experience, and practical knowledge of React, TypeScript, Tailwind, Python, SQL, and Vercel. Most people learning to code start from nothing.")
body("The gap between where you are and a professional developer role is:")
bullet("Depth, not breadth — go deeper in what you know before adding more")
bullet("Portfolio proof — built and deployed projects, not just tutorials")
bullet("Consistency — 30 hours a week for 5-6 months is genuinely enough")
bullet("Professional habits — Git workflow, testing, documentation, clean code")
tip("The most important thing is to build. Every day you write code is a day you improve. Every deployed project is evidence. Every line of tested, documented, clean code is the mark of a professional. You already have what it takes — now execute.")
body("Good luck. The developer job market needs people who can build things well. That is what you are becoming.")

doc.save('Professional_Developer_Roadmap.docx')
print("Done!")
